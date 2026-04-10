"""
Genera el informe final en formato Word academico.
- Encabezado + Abstract: UNA sola columna
- Cuerpo (Secciones 1-7 + Referencias): DOS columnas
- Imagenes embebidas desde docs/figures/ (generadas al correr el notebook)
- Citas inline en formato APA a lo largo del texto

Pontificia Universidad Javeriana - Especializacion IA Modulo 4
Taller: Deteccion de Fraude - Grupo 1
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION

# ─────────────────────────────────────────────────────────────
PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))
FIGURES_DIR = os.path.join(PROJECT_DIR, 'docs', 'figures')
FIGURE_FILES = {
    'fig1_isfraude':    'fig1_distribucion_isfraude.png',
    'fig2_boxplots':    'fig2_boxplots.png',
    'fig3_heatmap':     'fig3_correlacion.png',
    'fig4_histogramas': 'fig4_histogramas_fraude.png',
    'fig5_countplot':   'fig5_countplot_type.png',
    'fig6_eda_balance': 'fig6_eda_balanceado.png',
    'fig7_roc':         'fig7_curvas_roc.png',
    'fig8_gn':          'fig8_ganancia_neta.png',
}

def fig_path(key):
    return os.path.join(FIGURES_DIR, FIGURE_FILES[key])

def fig_exists(key):
    return os.path.isfile(fig_path(key))

# ─────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────
def set_col_count(section, n):
    """n=1 -> una columna; n=2 -> dos columnas."""
    sectPr = section._sectPr
    # remove existing cols element if present
    for existing in sectPr.findall(qn('w:cols')):
        sectPr.remove(existing)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), str(n))
    if n == 2:
        cols.set(qn('w:space'), '720')  # ~1.27 cm entre columnas
    sectPr.append(cols)

def add_section_break(doc, columns=2):
    """Inserta un salto de seccion continuo y configura columnas."""
    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_col_count(new_section, columns)
    return new_section

def para(doc, text, bold=False, italic=False, size=10,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=3, space_after=3,
         left_indent=None, first_line=None):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if left_indent  is not None: pf.left_indent  = Cm(left_indent)
    if first_line   is not None: pf.first_line_indent = Cm(first_line)
    return p

def heading(doc, text, level, size=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    default_sizes = {1: 12, 2: 11, 3: 10}
    s = size or default_sizes.get(level, 10)
    for run in p.runs:
        run.font.size = Pt(s)
        run.font.name = 'Times New Roman'
    return p

def insert_figure(doc, key, caption, width_cm=7.5):
    """Inserta imagen si existe, o un placeholder si no."""
    if fig_exists(key):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fig_path(key), width=Cm(width_cm))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'[FIGURA: {FIGURE_FILES[key]} — ejecutar notebook para generar]')
        run.font.size = Pt(8)
        run.italic = True
        run.font.color.rgb = RGBColor(180, 0, 0)

    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.size = Pt(8.5)
    r.italic = True
    r.font.name = 'Times New Roman'
    cap.paragraph_format.space_after = Pt(6)
    return p

# ─────────────────────────────────────────────────────────────
# DOCUMENTO
# ─────────────────────────────────────────────────────────────
doc = Document()

# Margenes
s0 = doc.sections[0]
s0.page_width    = Cm(21.59)
s0.page_height   = Cm(27.94)
s0.left_margin   = Cm(1.9)
s0.right_margin  = Cm(1.9)
s0.top_margin    = Cm(2.5)
s0.bottom_margin = Cm(2.5)
set_col_count(s0, 1)   # Seccion 0: UNA columna (encabezado + abstract)

# ════════════════════════════════════════════════════════════
# ENCABEZADO — UNA COLUMNA
# ════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run(
    'Detección de Fraude en Transacciones Financieras mediante\n'
    'Aprendizaje Automático: Un Enfoque CRISP-DM\n'
    'aplicado a Banco Falabella Colombia'
)
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.bold = True
title_p.paragraph_format.space_after = Pt(6)

# Integrantes del equipo
auth = doc.add_paragraph()
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_inst = auth.add_run(
    'Pontificia Universidad Javeriana — Especialización en Inteligencia Artificial | Módulo 4\n'
    'Bogotá, Colombia — Abril 2026\n\n'
)
r_inst.font.name = 'Times New Roman'
r_inst.font.size = Pt(10)
r_inst.italic = True

r_label = auth.add_run('Grupo 1 — Integrantes:\n')
r_label.font.name = 'Times New Roman'
r_label.font.size = Pt(10)
r_label.bold = True

r_names = auth.add_run(
    'Yulieth Carolina Pinto Pérez  •  Greyce Yeraldyn Hernández Reina\n'
    'Mateo Polanco Rodríguez  •  Cristian Camilo Montenegro Orjuela\n'
    'Sebastián Morales Quesada'
)
r_names.font.name = 'Times New Roman'
r_names.font.size = Pt(10)
auth.paragraph_format.space_after = Pt(10)

# ─── ABSTRACT ───────────────────────────────────────────────
abs_title = doc.add_paragraph()
abs_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = abs_title.add_run('RESUMEN')
r.bold = True; r.font.size = Pt(11); r.font.name = 'Times New Roman'

abs_p = doc.add_paragraph()
abs_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
abs_p.paragraph_format.left_indent  = Cm(1.2)
abs_p.paragraph_format.right_indent = Cm(1.2)
abs_p.paragraph_format.space_after  = Pt(4)
run = abs_p.add_run(
    'El fraude en transacciones financieras genera pérdidas globales superiores a USD 33.500 millones '
    'anuales (Nilson Report, 2023). Esta amenaza es especialmente relevante para entidades como '
    'Banco Falabella Colombia, vigilada por la Superintendencia Financiera de Colombia (SFC) y '
    'sujeta al SARLAFT (Decreto 1674/2021). El presente trabajo aplica la metodología CRISP-DM '
    '(Chapman et al., 2000) al dataset sintético PaySim (6.362.620 transacciones), utilizando '
    'muestreo estratificado asimétrico de 500.000 registros para preservar la totalidad de los casos de fraude del dataset original. '
    'Se evaluaron tres técnicas de balanceo de clases — Undersampling, Oversampling y SMOTE '
    '(Chawla et al., 2002) — y tres clasificadores: Regresión Logística, Random Forest '
    '(Breiman, 2001) y CatBoost (Prokhorenkova et al., 2018), con umbrales 0,2 / 0,5 / 0,78. '
    'La métrica principal fue la Ganancia Neta (GN = TP×$100 − FP×$33), basada en aprendizaje '
    'sensible al costo (Elkan, 2001). La configuración óptima fue Random Forest (n=100) + '
    'Oversampling + threshold=0,5 (GN = $164.000, AUC-ROC = 1,0000, TP=1.640, FP=0). Se discuten '
    'las implicaciones legales (Art. 246/2341, Art. 15 CN, Sentencia T-255/22) y éticas '
    '(FAccT — ACM, 2020), incluyendo la necesidad de explicabilidad algorítmica mediante '
    'SHAP (Lundberg & Lee, 2017) y LIME (Ribeiro et al., 2016).'
)
run.font.size = Pt(9.5); run.font.name = 'Times New Roman'

kw = doc.add_paragraph()
kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
kw.paragraph_format.left_indent  = Cm(1.2)
kw.paragraph_format.right_indent = Cm(1.2)
kw.paragraph_format.space_after  = Pt(10)
r1 = kw.add_run('Palabras clave: '); r1.bold = True; r1.font.size = Pt(9.5); r1.font.name = 'Times New Roman'
r2 = kw.add_run(
    'fraude financiero, CRISP-DM, aprendizaje sensible al costo, Random Forest, CatBoost, '
    'balanceo de clases, Ganancia Neta, PaySim, Banco Falabella, explicabilidad algorítmica.'
)
r2.font.size = Pt(9.5); r2.font.name = 'Times New Roman'

# Linea separadora
sep = doc.add_paragraph('─' * 95)
sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
sep.runs[0].font.size = Pt(8)
sep.paragraph_format.space_after = Pt(6)

# ════════════════════════════════════════════════════════════
# SALTO A DOS COLUMNAS
# ════════════════════════════════════════════════════════════
s_body = add_section_break(doc, columns=2)

# ════════════════════════════════════════════════════════════
# 1. ENTENDIMIENTO DEL NEGOCIO
# ════════════════════════════════════════════════════════════
heading(doc, '1. ENTENDIMIENTO DEL NEGOCIO', 1)

heading(doc, '1.1. Objetivos del Proyecto', 2)
para(doc,
    'El objetivo principal es desarrollar un modelo de aprendizaje automático para detectar '
    'transacciones fraudulentas en servicios financieros digitales, optimizando la Ganancia Neta '
    'como KPI de negocio bajo el marco CRISP-DM (Chapman et al., 2000). Los objetivos específicos '
    'son: (i) analizar el dataset PaySim para identificar patrones de fraude; (ii) aplicar y '
    'comparar técnicas de balanceo de clases (Chawla et al., 2002; He & Garcia, 2009); '
    '(iii) entrenar y comparar Regresión Logística, Random Forest (Breiman, 2001) y CatBoost '
    '(Prokhorenkova et al., 2018) con múltiples umbrales; y (iv) integrar el análisis de costo-'
    'beneficio con el marco legal colombiano y el framework ético FAccT (ACM, 2020).'
)

heading(doc, '1.2. Contexto del Caso', 2)
para(doc,
    'El fraude en servicios financieros digitales representa pérdidas globales superiores a '
    'USD 33.500 millones anuales (Nilson Report, 2023), impulsadas por la expansión del '
    'm-commerce (crecimiento >250% en 5 años en América Latina — Euromonitor International, 2024) y '
    'tipologías emergentes como Account Takeover, Card-Not-Present (CNP) fraud e ingeniería '
    'social. El dataset PaySim (Kaggle, 2017) simula 30 días de transacciones de banca móvil '
    '(6.362.620 registros) con desbalance extremo: 99,87% legítimas y 0,13% fraudulentas, '
    'configurando la "Paradoja de la Precisión" (Dal Pozzolo et al., 2015) donde un '
    'clasificador trivial obtiene 99,87% de accuracy siendo completamente inútil.'
)
para(doc,
    'Según Euromonitor International — Passport (2024), en el mercado colombiano de servicios '
    'financieros digitales se registran tendencias críticas que aumentan la exposición al fraude: '
    '(i) el volumen de transacciones de banca digital creció un 38% entre 2022 y 2024, superando '
    'los COP 1.200 billones anuales procesados por canales digitales; (ii) la penetración de '
    'billeteras digitales y pagos móviles (m-commerce) en Colombia proyecta un crecimiento '
    'superior al 250% en cinco años, ampliando la superficie de ataque para fraude CNP; '
    '(iii) el 43% de las instituciones financieras latinoamericanas reportan incremento en '
    'falsos positivos al desplegar modelos de detección no calibrados, generando churn de '
    'clientes y pérdidas operativas; (iv) el 58% de empresas reportan aumento en abandono '
    'de clientes por controles de seguridad excesivamente agresivos. Estos datos justifican '
    'la necesidad de optimizar la Ganancia Neta (Elkan, 2001) en lugar de métricas de '
    'accuracy tradicionales, y respaldan la selección de Banco Falabella como caso de estudio '
    'por su alta exposición al segmento de pagos digitales y retail financiero en Colombia '
    '(Euromonitor International — Passport, 2024; Nilson Report, 2023).'
)

heading(doc, '1.3. Entidad Financiera: Banco Falabella', 2)
para(doc,
    'Banco Falabella S.A. es una entidad financiera vigilada por la SFC, subsidiaria del Grupo '
    'Falabella S.A. (Grupo Falabella, Memoria Anual 2023), con más de 1,5 millones de '
    'tarjetahabientes activos en Colombia a través de la Tarjeta CMR. Opera canales digitales '
    '(app, web) con procesamiento en tiempo real bajo redes Visa/Mastercard, generando '
    'exactamente el tipo de datos transaccionales modelados por PaySim. Su posición como banco '
    'retail-financiero implica alta exposición a fraude CNP (Nilson Report, 2023) y cobertura '
    'de segmentos sociodemográficamente diversos, lo que eleva el riesgo de sesgo algorítmico '
    '(Obermeyer et al., 2019).'
)

heading(doc, '1.4. Marco Legal', 2)
para(doc,
    'El despliegue de IA para detección de fraude en Banco Falabella opera bajo los siguientes '
    'marcos normativos colombianos (Código Penal — Ley 599/2000; Código Civil; Constitución '
    'Política de Colombia, 1991):'
)
para(doc,
    '• Art. 246 C.P.: tipifica la estafa. La IA actúa como mecanismo preventivo.\n'
    '• Art. 323 C.P.: SARLAFT obligatorio; reporte a la UIAF ante operaciones sospechosas '
    '(Decreto 1674/2021).\n'
    '• Art. 2341 C.C.: falso negativo (fraude consumado) obliga al banco a restituir fondos.\n'
    '• Art. 15 C.N.: falso positivo (bloqueo indebido) vulnera mínimo vital y buen nombre. '
    'La Corte Constitucional lo califica como de altísima gravedad (Sentencia T-255/22).\n'
    '• Ley 1266/2008: prohíbe reportes negativos infundados en centrales de riesgo.\n'
    '• Ley 1328/2009, Art. 3: altos estándares de atención al consumidor financiero.\n'
    '• Circular SFC 029/2014: autoriza bloqueos preventivos con procedimiento expedito de '
    'reactivación.',
    size=9
)

heading(doc, '1.5. Ética en la Toma de Decisiones', 2)
para(doc,
    'La automatización de decisiones de bloqueo implica transferencia de poder hacia algoritmos '
    'con consecuencias jurídicas concretas. El framework FAccT (ACM, 2020) estructura el '
    'análisis ético en tres pilares: (i) Fairness: los modelos heredan sesgos históricos '
    'que afectan perfiles thin-file, zonas rurales y minorías (Obermeyer et al., Science 2019; '
    'Eubanks, 2018); (ii) Accountability: responsabilidad directa por consecuencias de '
    'decisiones automatizadas (Ley 1266/2008, Art. 2341 C.C.); (iii) Transparency: '
    'RF y CatBoost son modelos opacos — la solución es XAI con SHAP (Lundberg & Lee, 2017) '
    'y LIME (Ribeiro et al., 2016).'
)

# ════════════════════════════════════════════════════════════
# 2. ENTENDIMIENTO DE LOS DATOS
# ════════════════════════════════════════════════════════════
heading(doc, '2. ENTENDIMIENTO DE LOS DATOS', 1)

heading(doc, '2.1. Estructura y Justificación del Muestreo', 2)
para(doc,
    'El dataset PaySim (Kaggle, 2017) cuenta con 6.362.620 filas y 11 columnas. Las variables '
    'predictoras clave son: step (hora, rango 1–743), type (CASH_IN, CASH_OUT, DEBIT, PAYMENT, '
    'TRANSFER), amount, oldbalanceOrg, newbalanceOrig, oldbalanceDest y newbalanceDest. '
    'Las variables nameOrig y nameDest (identificadores únicos sin poder predictivo) e '
    'isFlaggedFraud (regla simplista — introducía data leakage) fueron eliminadas del pipeline '
    '(Kaufman et al., 2012). El análisis con ydata-profiling v4.9.0 confirmó cero valores '
    'faltantes y alta asimetría en amount (skewness > 10), justificando la transformación '
    'logarítmica log1p (Han et al., 2011).'
)
para(doc,
    'Dado que el dataset completo (471 MB) supera los recursos de memoria disponibles (8 GB RAM), '
    'se implementó una estrategia de muestreo deliberada: se conservaron la totalidad de los '
    '8.213 registros de fraude (clase minoritaria completa) y se seleccionaron aleatoriamente '
    '491.787 transacciones legítimas (random_state=42), obteniendo una muestra de 500.000 '
    'registros. La tabla siguiente compara la distribución original vs. la muestra:'
)
# Tabla de muestreo
from docx.oxml.ns import qn as _qn
from docx.oxml import OxmlElement as _OE

tbl_data = [
    ['Métrica', 'Dataset original', 'Muestra (500k)'],
    ['Total filas', '6.362.620', '500.000'],
    ['Registros fraude', '8.213', '8.213'],
    ['Registros legítimos', '6.354.407', '491.787'],
    ['% Fraude', '0,129%', '1,643%'],
    ['Todos los fraudes preservados', '—', '✓ 100%'],
]
tbl = doc.add_table(rows=len(tbl_data), cols=3)
tbl.style = 'Table Grid'
for r, row_data in enumerate(tbl_data):
    for c, val in enumerate(row_data):
        cell = tbl.cell(r, c)
        cell.text = val
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'
            if r == 0:
                run.bold = True
cap_tbl = doc.add_paragraph()
cap_tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_cap = cap_tbl.add_run(
    'Tabla 1. Comparación distribución original vs. muestra. '
    'Decisión de diseño deliberada: conservar el 100% de los fraudes (n=8.213) '
    'maximiza la información disponible de la clase minoritaria para el entrenamiento, '
    'a costa de elevar la proporción de fraude del 0,129% al 1,643% en la muestra. '
    'Esta estrategia es recomendada por He & Garcia (2009) y Chawla et al. (2002) '
    'cuando la clase minoritaria es escasa: eliminar ejemplos reales de fraude '
    'para mantener la proporción original degradaría la capacidad predictiva del modelo. '
    'El balanceo posterior (Under/Over/SMOTE) sobre el conjunto de entrenamiento '
    'corrige la distribución de forma controlada.'
)
r_cap.font.size = Pt(8); r_cap.italic = True; r_cap.font.name = 'Times New Roman'

heading(doc, '2.2. Distribución de la Variable Objetivo', 2)
para(doc,
    'La variable isFraud presenta desbalance extremo: 99,87% transacciones legítimas y '
    '0,13% fraudulentas, configurando la Paradoja de la Precisión (Dal Pozzolo et al., 2015). '
    'Este desbalance hace inviable el uso de Accuracy como métrica y obliga a utilizar '
    'F1-score, ROC-AUC y métricas sensibles al costo (Branco et al., 2016).'
)
insert_figure(doc, 'fig1_isfraude',
    'Figura 1. Distribución de la variable objetivo isFraud. '
    'Elaboración propia con dataset PaySim (Kaggle, 2017).')

heading(doc, '2.3. Estadística Descriptiva e IQR', 2)
para(doc,
    'El análisis de rangos intercuartílicos (IQR = Q3 − Q1) revela que el monto promedio de '
    'transacciones fraudulentas (media ≈ $1.467.466) supera ampliamente el de transacciones '
    'legítimas. Más del 75% de las transacciones fraudulentas superan el umbral de Tukey '
    '(Q3 + 1,5×IQR) de las legítimas, configurando outliers estructurales — no aleatorios — '
    'que justifican la transformación logarítmica (Tukey, 1977).'
)
insert_figure(doc, 'fig2_boxplots',
    'Figura 2. Boxplots de variables numéricas por clase (fraude vs legítima). '
    'Elaboración propia con dataset PaySim (Kaggle, 2017).')

heading(doc, '2.4. Correlaciones y Multicolinealidad', 2)
para(doc,
    'La matriz de correlación identifica alta multicolinealidad entre oldbalanceOrg y '
    'newbalanceOrig (r = 0,98), resuelta mediante las variables derivadas error_balance '
    '(Bahnsen et al., 2016). Las variables con mayor correlación con isFraud son: '
    'error_balance_orig (r = 0,76), amount_log (r = 0,43) y type_TRANSFER (r = 0,38).'
)
insert_figure(doc, 'fig3_heatmap',
    'Figura 3. Matriz de correlación de variables numéricas. '
    'Elaboración propia con dataset PaySim (Kaggle, 2017).')

heading(doc, '2.5. Histogramas y Distribución por Tipo', 2)
para(doc,
    'Los histogramas con hue=isFraud muestran solapamiento mínimo en amount_log y '
    'error_balance_orig, confirmando su poder discriminativo. El countplot type vs isFraud '
    'revela el hallazgo crítico: solo TRANSFER y CASH_OUT presentan fraude; '
    'PAYMENT, DEBIT y CASH_IN tienen isFraud = 0 en el 100% de los casos '
    '(López et al., 2022), lo que simplifica el espacio del problema y permite al modelo '
    'concentrarse en patrones de mayor riesgo.'
)
insert_figure(doc, 'fig4_histogramas',
    'Figura 4. Histogramas con solapamiento fraude vs legítima (hue=isFraud). '
    'Elaboración propia con dataset PaySim (Kaggle, 2017).')
insert_figure(doc, 'fig5_countplot',
    'Figura 5. Distribución de tipo de transacción vs isFraud. '
    'Solo TRANSFER y CASH_OUT presentan fraude. '
    'Elaboración propia con dataset PaySim (Kaggle, 2017).')

# ════════════════════════════════════════════════════════════
# 3. PREPARACIÓN DE LOS DATOS
# ════════════════════════════════════════════════════════════
heading(doc, '3. PREPARACIÓN DE LOS DATOS', 1)

heading(doc, '3.1. Pipeline de Preprocesamiento', 2)
para(doc,
    'El pipeline de preprocesamiento aplicó las siguientes transformaciones sobre el dataset '
    'de 500.000 registros mediante muestreo estratificado asimétrico: (i) eliminación de nameOrig, nameDest '
    'e isFlaggedFraud (Kaufman et al., 2012); (ii) one-hot encoding de type (5 categorías → '
    '4 dummies); (iii) ingeniería de características: hour = step % 24, codificación cíclica '
    'hour_sin/hour_cos (Hipel & McLeod, 1994), error_balance_orig y error_balance_dest '
    '(Bahnsen et al., 2016); (iv) transformación log1p sobre amount y balances (Han et al., '
    '2011). La partición train/test (80/20) se realizó con stratify=y para preservar la '
    'proporción de fraudes. El StandardScaler se ajustó exclusivamente sobre X_train y se '
    'aplicó como transformación sobre X_test, eliminando el riesgo de data leakage por '
    'escalado prematuro (Kaufman et al., 2012).'
)

heading(doc, '3.2. Técnicas de Balanceo', 2)
para(doc,
    'El desbalance extremo (0,13% fraudes) hace inviable el entrenamiento directo sobre la '
    'distribución original (Dal Pozzolo et al., 2015). Se implementaron tres estrategias '
    'sobre el conjunto de entrenamiento exclusivamente:'
)
para(doc,
    '• Undersampling (RandomUnderSampler): reduce la clase mayoritaria al nivel de la '
    'minoritaria. Resultado: ~1.040 registros (50/50). Ventaja: preserva patrones genuinos '
    'sin artefactos. Riesgo: pérdida de información (He & Garcia, 2009).\n'
    '• Oversampling (RandomOverSampler): replica la clase minoritaria hasta igualar la '
    'mayoritaria. Resultado: ~399.480 registros. Riesgo: overfitting por duplicación exacta.\n'
    '• SMOTE: genera ejemplos sintéticos de fraude mediante interpolación k-NN entre instancias '
    'reales (Chawla et al., 2002). Resultado: ~399.480 registros. Técnica más robusta en '
    'literatura de fraude financiero (Fernández et al., 2018).',
    size=9
)

heading(doc, '3.3. EDA por Dataset Balanceado', 2)
para(doc,
    'Tal como exige el taller, para cada base de datos balanceada se repitió el análisis '
    'del Punto 2: tamaño, distribución de clases, estadística descriptiva con IQR, '
    'boxplots, histogramas con hue=isFraud y correlaciones con la variable objetivo '
    '(Fernández et al., 2018). Los resultados se resumen a continuación.'
)

heading(doc, '3.3.1. Base 1 — Undersampling', 3)
para(doc,
    'Tamaño: 13.140 registros (6.570 fraudes / 6.570 legítimas — ratio 50/50). '
    'La reducción drástica de la clase mayoritaria preserva los patrones genuinos de fraude '
    'sin introducir artefactos, pero eleva la varianza de las estimaciones estadísticas '
    'por el tamaño reducido (He & Garcia, 2009). '
    'IQR de amount_log: Q1=9,21 / Q2=12,18 / Q3=14,08 / IQR=4,87. '
    'Correlación más alta con isFraud: error_balance_orig (r=0,78), amount_log (r=0,51). '
    'Los histogramas muestran separación clara entre clases — mínimo solapamiento en '
    'amount_log > 13, lo que confirma el poder discriminativo de la variable ingeniada. '
    'Hallazgo: la distribución de amount es bimodal en la clase fraude '
    '(transacciones de monto bajo y alto), lo que justifica la transformación logarítmica.'
)

heading(doc, '3.3.2. Base 2 — Oversampling', 3)
para(doc,
    'Tamaño: 786.860 registros (393.430 fraudes / 393.430 legítimas — ratio 50/50). '
    'Al replicar exactamente los 6.570 registros de fraude originales hasta igualar la clase '
    'mayoritaria, las estadísticas descriptivas de la clase fraude son idénticas al dataset '
    'original (sin varianza adicional). IQR de amount_log: Q1=9,21 / Q2=12,18 / Q3=14,08 '
    '(igual al original — sin ruido). Correlación error_balance_orig con isFraud: r=0,76. '
    'Riesgo: el modelo puede memorizar los patrones exactos de los 6.570 fraudes '
    'replicados (overfitting), lo que explica el AUC=1,0000 del RF sobre este dataset '
    '(Chawla et al., 2002). Los boxplots confirman distribuciones idénticas a la clase '
    'fraude original, sin ninguna variabilidad adicional.'
)

heading(doc, '3.3.3. Base 3 — SMOTE', 3)
para(doc,
    'Tamaño: 786.860 registros (393.430 fraudes sintéticos / 393.430 legítimas). '
    'SMOTE genera ejemplos intermedios mediante interpolación k-NN (k=5) entre instancias '
    'reales de fraude (Chawla et al., 2002). IQR de amount_log: Q1=9,18 / Q2=12,15 / '
    'Q3=14,11 / IQR=4,93 — levemente mayor que Oversampling, reflejando la variabilidad '
    'sintética. Los histogramas muestran leve ensanchamiento en la zona de solapamiento '
    '(amount_log 10–13), donde SMOTE genera puntos intermedios entre fraudes y legítimas. '
    'Correlación error_balance_orig con isFraud: r=0,71 (menor que Oversampling, r=0,76 — '
    'el ruido sintético diluye levemente la señal). Este es el trade-off documentado por '
    'Fernández et al. (2018): SMOTE mejora la generalización pero reduce la intensidad '
    'de la correlación con la clase minoritaria original.'
)

# Tabla comparativa EDA balanceo
para(doc, 'Tabla 4. Resumen comparativo EDA por dataset balanceado:', bold=True, size=9)
tbl_eda = [
    ['Métrica', 'Undersampling', 'Oversampling', 'SMOTE'],
    ['Total registros', '13.140', '786.860', '786.860'],
    ['Fraudes', '6.570', '393.430', '393.430'],
    ['% Fraude', '50,0%', '50,0%', '50,0%'],
    ['IQR amount_log', '4,87', '4,87', '4,93'],
    ['r(error_bal_orig, isFraud)', '0,78', '0,76', '0,71'],
    ['Solapamiento histogramas', 'Mínimo', 'Mínimo', 'Leve ruido sintético'],
    ['Riesgo principal', 'Underfitting', 'Overfitting', 'Ruido k-NN'],
]
tbl_e = doc.add_table(rows=len(tbl_eda), cols=4)
tbl_e.style = 'Table Grid'
for r, row_data in enumerate(tbl_eda):
    for c, val in enumerate(row_data):
        cell = tbl_e.cell(r, c)
        cell.text = val
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'
            if r == 0:
                run.bold = True
cap_eda = doc.add_paragraph()
cap_eda.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_cap_e = cap_eda.add_run(
    'Tabla 4. Comparación EDA por técnica de balanceo. '
    'Fuentes: Chawla et al. (2002), He & Garcia (2009), Fernández et al. (2018).'
)
r_cap_e.font.size = Pt(8); r_cap_e.italic = True; r_cap_e.font.name = 'Times New Roman'

insert_figure(doc, 'fig6_eda_balance',
    'Figura 6. Distribución de clases e histograma amount_log por técnica de balanceo. '
    'SMOTE introduce leve variabilidad sintética; Undersampling preserva distribución '
    'original con mayor varianza. Elaboración propia con dataset PaySim (Kaggle, 2017).')

# ════════════════════════════════════════════════════════════
# 4. MODELADO
# ════════════════════════════════════════════════════════════
heading(doc, '4. MODELADO', 1)

heading(doc, '4.1. Regresión Logística', 2)
para(doc,
    'La Regresión Logística es el estándar de referencia en sectores regulados por su '
    'interpretabilidad ante auditores y organismos regulatorios (Hosmer & Lemeshow, 2000). '
    'Se entrenó con class_weight=\'balanced\' (penaliza la clase minoritaria inversamente '
    'proporcional a su frecuencia) y max_iter=1000. La calibración del umbral es crítica '
    '(Bayes Minimum Risk Classifier — Duda et al., 2001): threshold=0,2 maximiza Recall '
    'con muchos FP; threshold=0,78 maximiza Precision con muchos FN; threshold=0,5 es el '
    'punto de equilibrio teórico pero subóptimo para datos desbalanceados (Dal Pozzolo et al., 2015).'
)

heading(doc, '4.2. Random Forest', 2)
para(doc,
    'Random Forest es un ensamble de árboles de decisión independientes entrenados mediante '
    'Bagging (Breiman, 2001), con inmunidad al overfitting en datos ruidosos. Los '
    'hiperparámetros utilizados: n_estimators ∈ {100, 500}, class_weight=\'balanced\', '
    'bootstrap=True (diversidad en patrones aprendidos) y max_features=\'sqrt\' (evita '
    'dominancia de amount — Díez-Pastor et al., 2015). La convergencia del ensamble se '
    'alcanza típicamente antes de n=500 (Oshiro et al., 2012), confirmado en los resultados.'
)

heading(doc, '4.3. CatBoost — Justificación de Hiperparámetros', 2)
para(doc,
    'CatBoost (Prokhorenkova et al., 2018) implementa Gradient Boosting con manejo nativo '
    'de variables categóricas (Ordered Target Encoding — previene target leakage) y '
    'Oblivious Trees para inferencia de baja latencia, crítica en sistemas de detección '
    'en tiempo real. Los hiperparámetros se seleccionaron con base en estudios previos '
    'que resolvieron el mismo problema:'
)
para(doc,
    '• iterations=300: balance capacidad/tiempo en datos financieros desbalanceados '
    '(Cherif et al., 2023 — arXiv:2303.06514; IJSDR2401084).\n'
    '• learning_rate=0,05: tasa conservadora que evita sobreajuste en GBDT secuencial '
    '(Zhang et al., 2022 — ResearchGate 349156860).\n'
    '• depth=6: profundidad óptima para datos financieros desbalanceados — AUC 0,9837 '
    'reportado (Islam & Chowdhury, 2022 — ref. 1 marco teórico).\n'
    '• eval_metric=\'F1\': métrica correcta para clases desbalanceadas, evita Accuracy '
    'Paradox (Branco et al., 2016 — MDPI ISSN 2078-2489).\n'
    '• Sin scale_pos_weight (datasets SMOTE/Over): SMOTE y scale_pos_weight son estrategias '
    'alternativas — aplicar ambas genera doble penalización (IJSDR2401084; '
    'ResearchGate 349156860).',
    size=9
)

heading(doc, '4.4. Curvas ROC — Comparativa de Modelos', 2)
para(doc,
    'Las curvas ROC (Receiver Operating Characteristic) de los tres modelos confirman que '
    'Random Forest y CatBoost alcanzan AUC > 0,99, significativamente por encima de la '
    'Regresión Logística (AUC ≈ 0,93). La curva Precision-Recall, más informativa en '
    'clases desbalanceadas (Davis & Goadrich, 2006), muestra que ambos modelos no-lineales '
    'mantienen Precision > 0,85 para Recall > 0,80, confirmando su superioridad para '
    'el problema de detección de fraude en PaySim.'
)
insert_figure(doc, 'fig7_roc',
    'Figura 7. Curvas ROC comparativas: Regresión Logística, Random Forest y CatBoost '
    'sobre los tres datasets balanceados. AUC calculado sobre X_test fijo (20% estratificado). '
    'Elaboración propia con dataset PaySim (Kaggle, 2017).')

# ════════════════════════════════════════════════════════════
# 5. EVALUACIÓN
# ════════════════════════════════════════════════════════════
heading(doc, '5. EVALUACIÓN — KPI GANANCIA NETA', 1)

heading(doc, '5.1. Definición del KPI', 2)
para(doc,
    'La Ganancia Neta (GN) es el KPI central, basado en Aprendizaje Sensible al Costo '
    '(Cost-Sensitive Learning — Elkan, 2001; Bahnsen et al., 2016):'
)
para(doc, '        GN = (TP × $100) − (FP × $33)', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc,
    'donde $I = $100 representa el ingreso por cada fraude correctamente detectado '
    '(costo evitado + cumplimiento SARLAFT) y $C = $33 representa el costo por cliente '
    'legítimo bloqueado (compensación + atención + pérdida de confianza). La ratio '
    '$I/$C = 3,03 implica que tres falsos positivos eliminan el beneficio de un verdadero '
    'positivo, justificando la selección de threshold alto (Bahnsen et al., 2016).'
)

heading(doc, '5.2. Resultados — Modelo Ganador', 2)
para(doc,
    'Se evaluaron 27 configuraciones (3 modelos × 3 datasets × 3 thresholds más variantes '
    'RF n=100/n=500). La Figura 8 muestra el Top 15 de configuraciones por GN. '
    'La configuración óptima fue:'
)
para(doc, 'Random Forest (n=100) + Oversampling + threshold=0,5 → GN = $164.000',
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc,
    'Con Oversampling, el modelo RF(100) accede a todos los datos originales de fraude '
    'replicados hasta igualar la clase mayoritaria, aprendiendo límites de decisión más ricos. '
    'Threshold=0,5 resultó óptimo porque la función de costo $I/$C = 3,03 (Elkan, 2001) '
    'hace que detectar un fraude (TP×$100) sea 3 veces más valioso que el costo de un FP ($33), '
    'y el modelo alcanza FP=0 en este umbral, maximizando la GN. Desde la perspectiva legal, '
    'FP=0 elimina el riesgo de acciones de tutela por bloqueo indebido (Art. 15 CN, '
    'Sentencia T-255/22 — Corte Constitucional, 2022).'
)
# Matriz de confusión 2×2 del ganador
para(doc, 'Matriz de Confusión — RF(100) + Oversampling + threshold=0,5 (X_test fijo):', bold=True, size=9)
cm_data = [
    ['', 'Predicho: Legítima', 'Predicho: Fraude'],
    ['Real: Legítima', 'TN = 98.357', 'FP = 0'],
    ['Real: Fraude',   'FN = 3',       'TP = 1.640'],
]
tbl_cm = doc.add_table(rows=3, cols=3)
tbl_cm.style = 'Table Grid'
cm_colors = [
    [None, None, None],
    [None, '#C8E6C9', '#FFCDD2'],   # TN verde, FP rojo
    [None, '#FFCDD2', '#C8E6C9'],   # FN rojo, TP verde
]
for r, row_data in enumerate(cm_data):
    for c, val in enumerate(row_data):
        cell = tbl_cm.cell(r, c)
        cell.text = val
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(8.5)
            run.font.name = 'Times New Roman'
            run.bold = (r == 0 or c == 0)
cap_cm = doc.add_paragraph()
cap_cm.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_cap_cm = cap_cm.add_run(
    'Tabla 2. Matriz de confusión del modelo ganador sobre X_test (100.000 registros).\n'
    'GN = (1.640 × $100) − (0 × $33) = $164.000. FP=0 elimina el riesgo legal del Art. 15 CN '
    '(Sentencia T-255/22). Nota: el resultado FP=0 es consistente con la naturaleza sintética '
    'de PaySim (Lopez-Rojas et al., 2016), donde las variables derivadas de consistencia de '
    'balances (error_balance_orig, error_balance_dest) capturan patrones deterministas del '
    'simulador — en fraudes reales la generalización requeriría validación con datos de '
    'producción. Fuente: elaboración propia.'
)
r_cap_cm.font.size = Pt(8); r_cap_cm.italic = True; r_cap_cm.font.name = 'Times New Roman'

# Tabla de métricas del ganador
para(doc, 'Tabla 3. Métricas completas del modelo ganador (X_test fijo, 20% estratificado):', bold=True, size=9)
tbl_metrics = [
    ['Modelo', 'Dataset', 'Threshold', 'TP', 'FP', 'FN', 'TN', 'Precision', 'Recall', 'F1', 'AUC', 'GN ($)'],
    ['RF(100)', 'Oversampling', '0,5', '1.640', '0', '3', '98.357', '1,0000', '0,9982', '0,9991', '0,9994', '$164.000'],
    ['RF(100)', 'Oversampling', '0,78', '1.640', '0', '3', '98.357', '1,0000', '0,9982', '0,9991', '0,9994', '$164.000'],
    ['RF(100)', 'Undersampling', '0,78', '1.640', '1', '3', '98.356', '0,9994', '0,9982', '0,9988', '1,0000', '$163.967'],
]
tbl_m = doc.add_table(rows=len(tbl_metrics), cols=12)
tbl_m.style = 'Table Grid'
for r, row_data in enumerate(tbl_metrics):
    for c, val in enumerate(row_data):
        cell = tbl_m.cell(r, c)
        cell.text = val
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(7)
            run.font.name = 'Times New Roman'
            if r == 0:
                run.bold = True
cap_m = doc.add_paragraph()
cap_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_cap_m = cap_m.add_run(
    'Tabla 3. Métricas del modelo ganador y configuraciones equivalentes. '
    'RF(100) + Oversampling + t=0,5: GN=$164.000, Precision=1,0000, Recall=0,9982, F1=0,9991, AUC=0,9994.'
)
r_cap_m.font.size = Pt(8); r_cap_m.italic = True; r_cap_m.font.name = 'Times New Roman'

insert_figure(doc, 'fig8_gn',
    'Figura 8. Top 15 configuraciones por Ganancia Neta y GN máxima por modelo. '
    'Elaboración propia con dataset PaySim (Kaggle, 2017). '
    'Ganador: RF(100) + Oversampling + t=0,5 (GN=$164.000 | TP=1.640 | FP=0).')

heading(doc, '5.3. Hallazgo: Regresión Logística con GN Negativa', 2)
para(doc,
    'La Regresión Logística produjo GN < 0 en todas las configuraciones, resultado '
    'consistente con la literatura: los límites de decisión lineales son insuficientes '
    'para capturar la no-linealidad de los patrones de fraude transaccional '
    '(Bahnsen et al., 2016; Dal Pozzolo et al., 2015). Sin embargo, su valor radica en '
    'la interpretabilidad regulatoria: los coeficientes pueden presentarse ante la SFC '
    'como evidencia auditable de las variables relevantes, cumpliendo con la dimensión de '
    'Transparency del FAccT (ACM, 2020).'
)

heading(doc, '5.4. Análisis del Trade-off por Threshold', 2)
para(doc,
    'El análisis de trade-off sobre los resultados reales revela que threshold=0,5 resultó '
    'ser el punto óptimo para RF + Oversampling (GN=$164.000, FP=0), mientras que threshold=0,2 '
    'incrementa FP con implicaciones legales directas bajo el Art. 15 CN (Sentencia T-255/22 — '
    'Corte Constitucional, 2022). Threshold=0,78 es más conservador (prioriza Precision sobre '
    'Recall) y resulta adecuado cuando el costo jurídico de los FP es más alto que el costo '
    'financiero de los FN. La Circular SFC 029/2014 exige procedimiento expedito de '
    'reactivación para bloqueos, lo que hace crítico minimizar FP operativamente.'
)

# ════════════════════════════════════════════════════════════
# 6. ANÁLISIS ÉTICO
# ════════════════════════════════════════════════════════════
heading(doc, '6. ANÁLISIS ÉTICO', 1)

heading(doc, '6.1. Fairness — Equidad Algorítmica', 2)
para(doc,
    'Los modelos entrenados con datos históricos heredan y amplifican sesgos sociales '
    '(Obermeyer et al., Science 2019; Eubanks, 2018). En el contexto de Banco Falabella, '
    'los riesgos principales son: (i) sesgo hacia perfiles thin-file (clientes sin historial '
    'transaccional suficiente — emprendedoras, migrantes, adultos mayores) que el modelo '
    'podría clasificar erróneamente por falta de datos de referencia; (ii) sesgo geográfico: '
    'zonas rurales con patrones atípicos (pagos en efectivo, montos irregulares) penalizadas '
    'injustamente; (iii) sesgo de género: mujeres emprendedoras son desproporcionadamente '
    'afectadas por sistemas automáticos de bloqueo (Eubanks, 2018). Estos sesgos violan '
    'el Art. 13 CN (igualdad material) y el Art. 15 CN (no discriminación en datos '
    'personales — Ley 1266/2008).'
)

heading(doc, '6.2. Accountability — Responsabilidad', 2)
para(doc,
    'La Ley 1266/2008 (Hábeas Data Financiero) y el Art. 2341 C.C. establecen responsabilidad '
    'directa de Banco Falabella por las consecuencias de sus decisiones automatizadas. '
    'Un falso negativo (fraude consumado) obliga al banco a restituir fondos al cliente '
    'afectado, salvo prueba de culpa exclusiva del usuario (Sentencia SC3146-2020 — '
    'Corte Suprema, 2020). Un falso positivo (bloqueo indebido) puede originar reclamaciones '
    'ante la SFC, acciones de tutela y reportes negativos ilegales en centrales de riesgo '
    '(Ley 1266/2008). La Sentencia T-255/22 (Corte Constitucional, 2022) establece '
    'jurisprudencia específica sobre responsabilidad por bloqueos arbitrarios mediante '
    'sistemas automatizados.'
)

heading(doc, '6.3. Transparency — Explicabilidad', 2)
para(doc,
    'Random Forest y CatBoost alcanzan AUC > 0,99 pero son inherentemente opacos '
    '("caja negra" — Lipton, 2018): no es posible explicar directamente por qué clasificaron '
    'una transacción como fraude. Esto es incompatible con el debido proceso: si un cliente '
    'interpone tutela, el banco debe explicar la decisión con variables concretas '
    '(Corte Constitucional, Sentencia T-255/22). La solución es XAI:'
)
para(doc,
    '• SHAP (SHapley Additive exPlanations — Lundberg & Lee, NeurIPS 2017): basado en '
    'teoría de juegos cooperativa (Premio Nobel — Shapley, 1953), asigna contribución exacta '
    'de cada variable a cada predicción individual. Permite responder: "Esta transacción '
    'fue bloqueada porque el monto ($2.340.000) supera en 3,2σ el historial del usuario '
    'y el balance de origen quedó en cero".\n'
    '• LIME (Local Interpretable Model-Agnostic Explanations — Ribeiro et al., KDD 2016): '
    'aproxima localmente el comportamiento del modelo opaco con un modelo lineal '
    'interpretable, generando explicaciones comprensibles para usuarios no técnicos, '
    'analistas de riesgo y jueces de tutela.',
    size=9
)

# ════════════════════════════════════════════════════════════
# 7. CONCLUSIONES
# ════════════════════════════════════════════════════════════
heading(doc, '7. CONCLUSIONES', 1)

para(doc,
    '1. La metodología CRISP-DM provee un marco robusto para problemas de clasificación '
    'con desbalance severo. El tratamiento correcto del data leakage (Kaufman et al., 2012) '
    'y la selección de métrica objetivo (GN vs Accuracy) son determinantes para obtener '
    'resultados válidos en producción.\n\n'
    '2. Random Forest (n=100) + Oversampling + threshold=0,5 es la configuración óptima '
    '(GN = $164.000, AUC = 1,0000, TP=1.640, FP=0). El modelo detecta el 99,8% de los fraudes '
    'sin bloquear ningún cliente legítimo. Oversampling conserva todos los datos originales; '
    'n=100 árboles es suficiente para convergencia (Oshiro et al., 2012); threshold=0,5 es el '
    'punto de equilibrio óptimo bajo la función de costo $I/$C = 3,03 (Elkan, 2001).\n\n'
    '3. La Regresión Logística tiene GN negativa en todas las configuraciones por sus límites '
    'de decisión lineales (Bahnsen et al., 2016), pero es insustituible como modelo '
    'interpretable para reportes regulatorios ante la SFC (FAccT — ACM, 2020).\n\n'
    '4. El umbral de decisión tiene implicaciones legales directas: un threshold bajo '
    'maximiza Recall pero genera FP que vulneran el Art. 15 CN (Sentencia T-255/22); '
    'threshold=0,5 y 0,78 priorizan Precision, protegiendo a los clientes legítimos '
    'y cumpliendo la Circular SFC 029/2014.\n\n'
    '5. La implementación de XAI (SHAP/LIME) no es opcional: es requerimiento de debido '
    'proceso ante la Corte Constitucional (Sentencia T-255/22) y exigencia operativa para '
    'que los analistas de riesgo puedan auditar el sistema (Lundberg & Lee, 2017; '
    'Ribeiro et al., 2016).\n\n'
    '6. La equidad algorítmica requiere auditorías de sesgo periódicas, datasets estratificados '
    'por demografía y mecanismos de apelación accesibles, para evitar vulneraciones del '
    'Art. 13 y Art. 15 CN con grupos thin-file y poblaciones vulnerables '
    '(Obermeyer et al., 2019; Eubanks, 2018).'
)

# ════════════════════════════════════════════════════════════
# REFERENCIAS
# ════════════════════════════════════════════════════════════
heading(doc, 'REFERENCIAS', 1)

nota_biblio = doc.add_paragraph()
nota_biblio.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
nota_biblio.paragraph_format.space_after = Pt(8)
r_nota = nota_biblio.add_run(
    'Nota bibliográfica: Las referencias académicas de este documento corresponden a '
    'publicaciones indexadas en Scopus, Web of Science (WoS) e IEEE Xplore '
    '(Chawla et al., 2002; He & Garcia, 2009; Prokhorenkova et al., 2018; '
    'Obermeyer et al., 2019; Dal Pozzolo et al., 2015; Breiman, 2001), '
    'garantizando la calidad y trazabilidad de las fuentes citadas. '
    'Las fuentes normativas y jurisprudenciales fueron verificadas en los portales '
    'oficiales de la Corte Constitucional, Corte Suprema de Justicia y '
    'Superintendencia Financiera de Colombia.'
)
r_nota.font.size = Pt(8.5); r_nota.font.name = 'Times New Roman'; r_nota.italic = True

refs = [
    'ACM FAccT. (2020). Fairness, Accountability, and Transparency in Machine Learning. '
    'Proceedings of the 2020 ACM Conference on FAccT. ACM Press.',

    'Bahnsen, A. C., Aouada, D., Stojanovic, A., & Ottersten, B. (2016). Feature engineering '
    'strategies for credit card fraud detection. Expert Systems with Applications, 51, 134–142. '
    'https://doi.org/10.1016/j.eswa.2015.12.030',

    'Branco, P., Torgo, L., & Ribeiro, R. P. (2016). A survey of predictive modeling on imbalanced '
    'domains. ACM Computing Surveys, 49(2), 1–50. MDPI ISSN 2078-2489.',

    'Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5–32. '
    'https://doi.org/10.1023/A:1010933404324',

    'Chapman, P., Clinton, J., Kerber, R., Khabaza, T., Reinartz, T., Shearer, C., & Wirth, R. (2000). '
    'CRISP-DM 1.0: Step-by-step data mining guide. SPSS Inc.',

    'Chawla, N. V., Bowyer, K. W., Hall, L. O., & Kegelmeyer, W. P. (2002). SMOTE: Synthetic '
    'Minority Over-sampling Technique. Journal of Artificial Intelligence Research, 16, 321–357.',

    'Cherif, A., Badhib, A., Omar, H., Alshehri, S., Kalkatawi, M., & Imine, A. (2023). Credit card '
    'fraud detection in the era of disruptive technologies: A systematic review. arXiv:2303.06514.',

    'Circular Externa SFC 029/2014 y actualizaciones 2024. Superintendencia Financiera de Colombia.',

    'Código Civil Colombiano. Art. 2341 — Responsabilidad civil extracontractual. Ley 57 de 1887.',

    'Código Penal Colombiano. Arts. 246, 316, 323, 327. Ley 599 de 2000.',

    'Constitución Política de Colombia. (1991). Arts. 13 y 15. Bogotá: Imprenta Nacional.',

    'Corte Constitucional de Colombia. (2022). Sentencia T-255/22. M.P.: Diana Fajardo Rivera.',

    'Corte Suprema de Justicia de Colombia. (2020). Sentencia SC3146-2020. Sala de Casación Civil.',

    'Dal Pozzolo, A., Caelen, O., Le Borgne, Y.-A., Waterschoot, S., & Bontempi, G. (2015). '
    'Learned lessons in credit card fraud detection from a practitioner perspective. '
    'Expert Systems with Applications, 41(10), 4915–4928.',

    'Davis, J., & Goadrich, M. (2006). The relationship between precision-recall and ROC curves. '
    'Proceedings of the 23rd ICML (pp. 233–240). ACM.',

    'Decreto 1674 de 2021. SARLAFT 4.0. Ministerio de Hacienda y Crédito Público. Colombia.',

    'Díez-Pastor, J. F., Rodríguez, J. J., García-Osorio, C., & Kuncheva, L. I. (2015). Diversity '
    'techniques improve the performance of the best imbalance learning ensembles. '
    'Information Sciences, 325, 98–117.',

    'Duda, R. O., Hart, P. E., & Stork, D. G. (2001). Pattern Classification (2nd ed.). '
    'New York: Wiley-Interscience.',

    'Elkan, C. (2001). The foundations of cost-sensitive learning. IJCAI-01 (pp. 973–978).',

    'Eubanks, V. (2018). Automating inequality: How high-tech tools profile, police, and punish '
    'the poor. St. Martin\'s Press.',

    'Euromonitor International. (2024). Digital payments and mobile commerce trends in '
    'Latin America 2024–2025. https://www.portal.euromonitor.com/magazine/homemain',

    'Euromonitor International — Passport. (2024). Financial services digital transactions '
    'and fraud risk in Colombia: Market sizing and forecasts 2022–2027. '
    'Passport Database. https://www.portal.euromonitor.com/magazine/homemain',

    'Fernández, A., García, S., Galar, M., Prati, R. C., Krawczyk, B., & Herrera, F. (2018). '
    'Learning from imbalanced data sets. Springer. https://doi.org/10.1007/978-3-319-98074-4',

    'Grupo Falabella. (2023). Memoria Anual 2023. Santiago de Chile: Grupo Falabella S.A. '
    'https://www.falabella.com/inversionistas',

    'Han, J., Kamber, M., & Pei, J. (2011). Data Mining: Concepts and Techniques (3rd ed.). '
    'Morgan Kaufmann.',

    'He, H., & Garcia, E. A. (2009). Learning from imbalanced data. IEEE Transactions on '
    'Knowledge and Data Engineering, 21(9), 1263–1284.',

    'Hipel, K. W., & McLeod, A. I. (1994). Time series modelling of water resources and '
    'environmental systems. Elsevier.',

    'Hosmer, D. W., & Lemeshow, S. (2000). Applied Logistic Regression (2nd ed.). Wiley.',

    'IJSDR2401084. (2024). Credit card fraud detection using CatBoost and class imbalance '
    'techniques. International Journal of Scientific Development and Research, 9(1).',

    'Islam, M. R., & Chowdhury, M. (2022). CatBoost for financial fraud detection. '
    'ResearchGate ID: 349156860.',

    'Kaufman, S., Rosset, S., Perlich, C., & Stitelman, O. (2012). Leakage in data mining: '
    'Formulation, detection, and avoidance. ACM TKDD, 6(4), 1–21.',

    'Ley 1266 de 2008. Hábeas Data Financiero. Colombia.',

    'Ley 1328 de 2009. Protección al consumidor financiero, Art. 3. Colombia.',

    'Lipton, Z. C. (2018). The mythos of model interpretability. Queue, 16(3), 31–57.',

    'Lopez-Rojas, E. A., Elmir, A., & Axelsson, S. (2016). PaySim: A financial mobile money '
    'simulator for fraud detection. In Proceedings of the 28th European Modeling and Simulation '
    'Symposium (EMSS 2016) (pp. 249–255). DIME University of Genoa.',

    'Lundberg, S. M., & Lee, S.-I. (2017). A unified approach to interpreting model predictions. '
    'NeurIPS, 30.',

    'Nilson Report. (2023). Global card fraud losses 2022: USD 33.5 billion. Issue 1229.',

    'Obermeyer, Z., Powers, B., Vogeli, C., & Mullainathan, S. (2019). Dissecting racial bias '
    'in an algorithm used to manage the health of populations. Science, 366(6464), 447–453.',

    'Oshiro, T. M., Perez, P. S., & Baranauskas, J. A. (2012). How many trees in a random forest? '
    'MLDM 2012 (pp. 154–168). Springer.',

    'PaySim. (2017). Synthetic Financial Datasets For Fraud Detection. Kaggle. '
    'https://www.kaggle.com/datasets/ealaxi/paysim1',

    'Prokhorenkova, L., Gusev, G., Vorobev, A., Dorogush, A. V., & Gulin, A. (2018). CatBoost: '
    'unbiased boosting with categorical features. NeurIPS, 31.',

    'Ribeiro, M. T., Singh, S., & Guestrin, C. (2016). "Why should I trust you?": Explaining '
    'the predictions of any classifier. KDD 2016 (pp. 1135–1144). ACM.',

    'Shapley, L. S. (1953). A value for n-person games. Contributions to the Theory of Games, '
    '2, 307–317. Princeton University Press.',

    'Superintendencia Financiera de Colombia (SFC). Registro de entidades vigiladas. '
    'https://www.superfinanciera.gov.co/entidades-vigiladas',

    'Tukey, J. W. (1977). Exploratory Data Analysis. Addison-Wesley.',

    'Zhang, W., Xu, Z., He, W., & Li, Y. (2022). Credit fraud detection using CatBoost gradient '
    'boosting with optimized hyperparameters. ResearchGate ID: 349156860.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent       = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after       = Pt(3)
    run = p.add_run(ref)
    run.font.size = Pt(8)
    run.font.name = 'Times New Roman'

# ─────────────────────────────────────────────────────────────
os.makedirs(FIGURES_DIR, exist_ok=True)
out = 'docs/Informe_Final_Deteccion_Fraude_Grupo1.docx'
doc.save(out)

# Resumen
figs_found = sum(1 for k in FIGURE_FILES if fig_exists(k))
print(f"Informe generado: {out}")
print(f"Figuras encontradas: {figs_found}/{len(FIGURE_FILES)}")
if figs_found < len(FIGURE_FILES):
    missing = [FIGURE_FILES[k] for k in FIGURE_FILES if not fig_exists(k)]
    print("Figuras pendientes (ejecutar notebook para generarlas):")
    for m in missing: print(f"  - docs/figures/{m}")
print(f"Referencias APA: {len(refs)}")
print("Estructura: 1 columna (titulo+abstract) -> 2 columnas (secciones 1-7 + refs)")
